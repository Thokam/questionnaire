import pprint
import random
import sys
from docx import Document
import argparse

#special rules
 
 
def anzahl(thema):
    anzahl = 0

    if thema == 0:
        anzahl = 2
    elif thema == 1:
        anzahl = 2
    elif thema == 2:
        anzahl = 2
    elif thema == 3:
        anzahl = 1
    elif thema == 4:
        anzahl = 1
    elif thema == 5:
        anzahl = 5
    elif thema == 6:
        anzahl = 5
    elif thema == 7:
        anzahl = 5
    elif thema == 8:
        anzahl = 5
    elif thema == 9:
        anzahl = 5
    elif thema == 10:
        anzahl = 5
    elif thema == 11:
        anzahl = 2
    elif thema == 12:
        anzahl = 5
    elif thema == 13:
        anzahl = 5
    else:
        anzahl = 0
    return anzahl


def get_questions(filename):
    temp_fragen=[]
    thema_fragen=[]
    with open(filename, "r") as file:
        
        for line in file:
            line.strip()
            if line != "###\n":
                thema_fragen.append(line)
            else:
                temp_fragen.append(thema_fragen)
                thema_fragen = []
                #feld +=1
        temp_fragen.append(thema_fragen)
    return temp_fragen

def get_hard_questions(filename_hard):
    temp_fragen = get_questions(filename_hard)
    result=[]
    used = 0
    used_qpt = []

    for x in range(0, len(temp_fragen)):
        used_qpt.append(0)
    while used <args.number_qh:
        t = random.randint(0, len(temp_fragen)-1)
        if used_qpt[t] < anzahl(t):
            output = random.sample(temp_fragen[t], 1)
            if result.count(output[0])==0:
                result.append(output[0])
                used +=1
                used_qpt[t] +=1  

    return result, used_qpt

def create_fragebogen():

    #All questions. One sublist per topic. Topics are seperated by '###'
    all_questions = []

    #used questions per topic 
    used_qpt = {}
    used_q=0

    #random id. Unique for questionnaire (as text and docx) and answerfile  
    id = random.randint(1,100000000)

    all_questions = get_questions(args.filename)
  
    #subset of all_questions
    random_questions = []

    
    if args.filename_hard != "0":
        random_questions,used_qpt = get_hard_questions(args.filename_hard)

    for topic in range(0, len(all_questions)):
        if args.number_q != -1:
            number_questions=args.number_q
        else: 
            number_questions = anzahl(topic)-used_qpt[topic]

        if len(all_questions[topic]) > 0:
            output = random.sample(all_questions[topic], number_questions)
            for item in output:
                random_questions.append(item)       

    #debug-output
    #pprint.pprint(random_questions)


    #create questionnaire and a file named 'check' with the correct answers in the correct order 
    with open(args.output_file+"_"+str(id)+".txt", "a") as out:
        with open (args.output_file+"_check_"+str(id)+".txt", "w") as control:
            q = random.sample(random_questions, len(random_questions))
            for item in q:
                f = item.split("?")[0]
                out.write(f+"?\n")
                antworten = item.split("?")[1].split("#")[0].split(",")
                for a in antworten:
                    out.write("O "+a+"\n")
                
                out.write("\n")
                control.write(f+": "+(str)(item.split("?")[1].split("#")[1]))

    #TODO Switch to docxtpl and use a template?
    #copy questionnaire to a docx-file
    with open(args.output_file+"_"+str(id)+".txt", "r") as input:
        content =input.read()

        document = Document()

        document.add_heading('Fragebogen: '+str(id), 0)
        document.add_paragraph(content)


        document.save(args.output_file+'_'+str(id)+'.docx')

if __name__ == "__main__":

    parser = argparse.ArgumentParser(description='Create questionnaires')
    parser.add_argument('-c', '--count', type=int, default=1, help='number of questionnaires to create. Default = 1')
    parser.add_argument('-f', '--filename', type=str, default="fragen.txt", help='file with the questions. Default=fragen.txt')
    parser.add_argument('-n', '--number_q', type=int, default=1, help='number of questions per topic. Default=1. Use -1 for the special rules')
    parser.add_argument('-o', '--output_file', type=str, default="antwort", help='filename-prefix for the questionnaire. Default=antwort')
    parser.add_argument('-qh', '--number_qh', type=int, default=0, help='Number of hard questions. Default = 0')
    parser.add_argument('-fh', '--filename_hard', type=str, default="0", help='file with the hard questions.')
    

    args=parser.parse_args()

    for x in range(0,args.count):
         create_fragebogen()